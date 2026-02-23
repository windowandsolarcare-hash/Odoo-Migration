import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_manual():
    # Read the extracted content
    with open("extracted_content.txt", "r", encoding="utf-8") as f:
        lines = f.readlines()
    
    doc = Document()
    
    # Global Style Adjustments
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    def add_heading(text, level):
        h = doc.add_heading(text, level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return h
    
    def add_section_separator():
        """Add a visual separator between sections"""
        p = doc.add_paragraph()
        p.add_run("=" * 80)
        p.runs[0].font.color.rgb = RGBColor(192, 192, 192)

    # ===== TITLE PAGE =====
    title = doc.add_heading('Window & Solar Care: Complete System Integration Manual', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Version: 5.0 (Master Revision - Comprehensive Edition)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: Sunday, Feb 1, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Generated from: Copy of System Inst to review.docx').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ===== TABLE OF CONTENTS PLACEHOLDER =====
    add_heading('Table of Contents', 1)
    toc_items = [
        "I. Workiz API Integration Protocol",
        "II. Odoo API Integration Protocol", 
        "III. System Architecture & Golden Rules",
        "IV. Master System Credentials",
        "V. Odoo Technical Constraints",
        "VI. Complete Field Dictionary - Custom Fields",
        "VII. Complete Field Dictionary - Base Fields",
        "VIII. Integration Workflows"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

    # ===== PARSE CONTENT INTO SECTIONS =====
    current_section = []
    sections = {}
    current_key = "intro"
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Detect major section breaks
        if line.startswith("*****"):
            if current_section:
                sections[current_key] = current_section
                current_section = []
            current_key = f"section_{len(sections)}"
        else:
            current_section.append(line)
    
    # Add final section
    if current_section:
        sections[current_key] = current_section

    # ===== SECTION I: WORKIZ API PROTOCOL =====
    add_heading('I. Workiz API Integration Protocol (Complete)', 1)
    
    # Find Workiz content
    workiz_content = []
    for key, content in sections.items():
        first_lines = ' '.join(content[:5]).lower()
        if 'workiz api' in first_lines or 'workiz' in content[0].lower():
            workiz_content.extend(content)
            if len(workiz_content) > 50:  # Get substantial content
                break
    
    # Write Workiz protocol
    in_list = False
    for line in workiz_content[:100]:  # First 100 lines of Workiz content
        if line.startswith(("1.", "2.", "3.", "4.", "·", "•", "-")):
            doc.add_paragraph(line, style='List Bullet')
            in_list = True
        elif ":" in line and len(line) < 100 and not line.startswith(" "):
            add_heading(line, 3)
            in_list = False
        else:
            if in_list and line and not line[0].isdigit():
                doc.add_paragraph(line, style='List Bullet 2')
            else:
                doc.add_paragraph(line)
                in_list = False
    
    add_section_separator()
    doc.add_page_break()

    # ===== SECTION II: ODOO API PROTOCOL =====
    add_heading('II. Odoo API Integration Protocol (Complete)', 1)
    
    # Find Odoo protocol content
    odoo_content = []
    for key, content in sections.items():
        first_lines = ' '.join(content[:5]).lower()
        if 'odoo api' in first_lines and 'protocol' in first_lines:
            odoo_content.extend(content)
            if len(odoo_content) > 50:
                break
    
    # Write Odoo protocol
    in_list = False
    for line in odoo_content[:150]:  # First 150 lines
        if line.startswith(("1.", "2.", "3.", "4.", "·", "•", "-")):
            doc.add_paragraph(line, style='List Bullet')
            in_list = True
        elif ":" in line and len(line) < 100 and not line.startswith(" "):
            add_heading(line, 3)
            in_list = False
        else:
            if in_list and line and len(line) > 10:
                doc.add_paragraph(line, style='List Bullet 2')
            else:
                doc.add_paragraph(line)
                in_list = False
    
    add_section_separator()
    doc.add_page_break()

    # ===== SECTION III: GOLDEN RULES =====
    add_heading('III. System Architecture & Golden Rules', 1)
    
    # Find Golden Rules content
    golden_rules = []
    for key, content in sections.items():
        for i, line in enumerate(content):
            if 'golden rule' in line.lower() or 'architecture law' in line.lower():
                golden_rules.extend(content[max(0, i-5):i+20])
                break
        if golden_rules:
            break
    
    if golden_rules:
        for line in golden_rules:
            if line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "·", "•")):
                doc.add_paragraph(line, style='List Bullet')
            elif ":" in line and len(line) < 100:
                p = doc.add_paragraph(line)
                p.runs[0].bold = True
            else:
                doc.add_paragraph(line)
    
    add_section_separator()
    doc.add_page_break()

    # ===== SECTION IV: CREDENTIALS =====
    add_heading('IV. Master System Credentials', 1)
    doc.add_paragraph('WARNING: SENSITIVE INFORMATION - Handle with care')
    
    creds = [
        ["System", "Key Type", "Value"],
        ["Odoo", "API Key", "7e92006fd5c71e4fab97261d834f2e6004b61dc6"],
        ["Workiz", "API Token", "api_1hu6lroiy5zxomcpptuwsg8heju97iwg"],
        ["Workiz", "API Secret", "sec_334084295850678330105471548"]
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(creds[0]):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    for sys, k_type, val in creds[1:]:
        row_cells = table.add_row().cells
        row_cells[0].text = sys
        row_cells[1].text = k_type
        row_cells[2].text = val
    
    add_section_separator()
    doc.add_page_break()

    # ===== SECTION V: ODOO CONSTRAINTS =====
    add_heading('V. Odoo Development Constraints & Sandbox Rules', 1)
    
    constraints = []
    for key, content in sections.items():
        for i, line in enumerate(content):
            if 'sandbox' in line.lower() or 'constraint' in line.lower():
                constraints.extend(content[i:i+30])
                break
        if constraints:
            break
    
    if constraints:
        for line in constraints[:40]:
            if line.startswith(("1.", "2.", "3.", "·", "•", "-")):
                doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line)
    
    add_section_separator()
    doc.add_page_break()

    # ===== SECTION VI: FIELD DICTIONARY - CUSTOM FIELDS =====
    add_heading('VI. Complete Field Dictionary - Custom Fields', 1)
    doc.add_paragraph('This section contains ALL custom fields extracted from the system.')
    
    # Parse custom fields from extracted content
    custom_fields = []
    in_custom_section = False
    
    for line in lines:
        line = line.strip()
        # Detect custom field lines (they contain x_studio or x_)
        if line.startswith("x_") or line.startswith('"x_'):
            parts = line.split('\t')
            if len(parts) >= 4:
                custom_fields.append(parts)
    
    if custom_fields:
        # Create table for custom fields
        add_heading('Custom Fields Table', 2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        headers = ["Field Name", "Label", "Model", "Type", "Stored"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Data rows (limit to first 100 for readability)
        for field_data in custom_fields[:100]:
            row_cells = table.add_row().cells
            for i in range(min(5, len(field_data))):
                row_cells[i].text = field_data[i]
        
        if len(custom_fields) > 100:
            doc.add_paragraph(f"Note: Showing first 100 of {len(custom_fields)} custom fields.")
    
    add_section_separator()
    doc.add_page_break()

    # ===== SECTION VII: BASE FIELDS =====
    add_heading('VII. Complete Field Dictionary - Base Fields', 1)
    doc.add_paragraph('This section contains ALL base Odoo fields.')
    
    # Parse base fields (CSV format lines with "Base Field")
    base_fields = []
    for line in lines:
        line = line.strip()
        if '"Base Field"' in line or 'Base Field' in line:
            # Clean up and split
            parts = line.replace('"', '').split(',')
            if len(parts) >= 4:
                base_fields.append(parts)
    
    if base_fields:
        add_heading('Sales Order Base Fields', 2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        headers = ["Field Name", "Label", "Model", "Type", "Indexed"]
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add rows (limit to first 50)
        for field_data in base_fields[:50]:
            row_cells = table.add_row().cells
            for i in range(min(5, len(field_data))):
                row_cells[i].text = field_data[i]
        
        if len(base_fields) > 50:
            doc.add_paragraph(f"Note: Showing first 50 of {len(base_fields)} base fields.")
    
    add_section_separator()
    doc.add_page_break()

    # ===== SECTION VIII: INTEGRATION WORKFLOWS =====
    add_heading('VIII. Key Integration Workflows', 1)
    
    workflow_content = []
    for key, content in sections.items():
        for i, line in enumerate(content):
            if 'lever pull' in line.lower() or 'two-step' in line.lower():
                workflow_content.extend(content[i:i+15])
                break
        if workflow_content:
            break
    
    add_heading('The "Lever Pull" SMS Workflow', 2)
    if workflow_content:
        for line in workflow_content[:20]:
            doc.add_paragraph(line, style='List Bullet')
    
    add_section_separator()

    # ===== APPENDIX =====
    doc.add_page_break()
    add_heading('Appendix: Complete Raw Content', 1)
    doc.add_paragraph('The following contains all extracted content for reference:')
    
    # Add all content in small font
    for line in lines[:500]:  # First 500 lines to keep document reasonable
        p = doc.add_paragraph(line.strip())
        p.runs[0].font.size = Pt(8)
    
    if len(lines) > 500:
        doc.add_paragraph(f"... (truncated, showing 500 of {len(lines)} total lines)")

    # ===== SAVE DOCUMENT =====
    doc_path = "Window_Solar_Care_Complete_Manual.docx"
    doc.save(doc_path)
    print(f"Manual created successfully!")
    print(f"File saved as: {doc_path}")
    print(f"Document contains {len(lines)} lines of content")
    print(f"Custom fields found: {len(custom_fields)}")
    print(f"Base fields found: {len(base_fields)}")

if __name__ == "__main__":
    create_manual()
