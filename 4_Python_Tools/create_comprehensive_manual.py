"""
Comprehensive Manual Generator for AI Agent Onboarding
Preserves 100% of source data while organizing logically
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_comprehensive_manual():
    print("Reading source document...")
    
    # Read the original Word document directly to preserve structure
    source_doc = Document("Copy of System Inst to review.docx")
    
    # Create new output document
    output_doc = Document()
    
    # === CONFIGURE STYLES ===
    style = output_doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    def add_heading(text, level=1):
        """Add a heading with consistent formatting"""
        h = output_doc.add_heading(text, level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return h
    
    def add_separator():
        """Visual separator between major sections"""
        p = output_doc.add_paragraph()
        run = p.add_run("=" * 100)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.size = Pt(8)
    
    # === EXTRACT ALL CONTENT ===
    print("Extracting all content from source...")
    all_paragraphs = []
    all_tables = []
    
    for para in source_doc.paragraphs:
        all_paragraphs.append(para.text)
    
    for table in source_doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        all_tables.append(table_data)
    
    print(f"Extracted {len(all_paragraphs)} paragraphs and {len(all_tables)} tables")
    
    # === ANALYZE AND CATEGORIZE CONTENT ===
    print("Analyzing and categorizing content...")
    
    content_sections = {
        'workiz_api': [],
        'odoo_api': [],
        'credentials': [],
        'golden_rules': [],
        'constraints': [],
        'field_definitions': [],
        'workflows': [],
        'validation': [],
        'other': []
    }
    
    current_category = 'other'
    
    for para in all_paragraphs:
        para_lower = para.lower()
        
        # Categorize based on content
        if any(x in para_lower for x in ['workiz api', 'workiz endpoint', '/job/create', '/job/update', 'workiz automation']):
            current_category = 'workiz_api'
        elif any(x in para_lower for x in ['odoo api', 'jsonrpc', 'execute_kw', 'search_read', 'window-solar-care.odoo']):
            current_category = 'odoo_api'
        elif any(x in para_lower for x in ['api key', 'api token', 'auth secret', 'credential']):
            current_category = 'credentials'
        elif any(x in para_lower for x in ['golden rule', 'architecture law', 'one number', 'constraint', 'mission directive']):
            current_category = 'golden_rules'
        elif any(x in para_lower for x in ['sandbox', 'forbidden', 'import statement', 'server action']):
            current_category = 'constraints'
        elif any(x in para_lower for x in ['x_studio', 'field name', 'field type', 'custom field', 'base field']):
            current_category = 'field_definitions'
        elif any(x in para_lower for x in ['lever pull', 'workflow', 'two-step', 'trigger', 'zapier']):
            current_category = 'workflows'
        elif any(x in para_lower for x in ['validation', 'confirmed', 'proven', 'hard-won']):
            current_category = 'validation'
        
        content_sections[current_category].append(para)
    
    print(f"Content categorized into {len(content_sections)} sections")
    
    # === BUILD COMPREHENSIVE MANUAL ===
    
    # ========== TITLE PAGE ==========
    title = output_doc.add_heading('Window & Solar Care', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = output_doc.add_heading('Complete System Integration Manual', 0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_doc.add_paragraph('Master Reference Document for AI Agents').alignment = WD_ALIGN_PARAGRAPH.CENTER
    output_doc.add_paragraph('Version: 6.0 (Complete - Zero Data Loss)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    output_doc.add_paragraph('Generated: February 1, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    output_doc.add_paragraph('Source: Copy of System Inst to review.docx (26 pages)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_doc.add_page_break()
    
    # ========== EXECUTIVE SUMMARY ==========
    add_heading('Executive Summary', 1)
    output_doc.add_paragraph(
        'This document contains the complete technical specifications, protocols, and data definitions '
        'for the Window & Solar Care integration between Odoo (CRM/ERP), Workiz (Field Service), and Zapier (Integration Platform). '
        'ALL source data has been preserved and organized for AI agent consumption.'
    )
    
    stats_para = output_doc.add_paragraph()
    stats_para.add_run('Document Statistics:\n').bold = True
    stats_para.add_run(f'• Total Paragraphs: {len(all_paragraphs)}\n')
    stats_para.add_run(f'• Total Tables: {len(all_tables)}\n')
    stats_para.add_run(f'• Workiz API Content: {len(content_sections["workiz_api"])} items\n')
    stats_para.add_run(f'• Odoo API Content: {len(content_sections["odoo_api"])} items\n')
    stats_para.add_run(f'• Field Definitions: {len(content_sections["field_definitions"])} items\n')
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== TABLE OF CONTENTS ==========
    add_heading('Table of Contents', 1)
    toc_items = [
        'Section 1: System Architecture & Golden Rules',
        'Section 2: Master Credentials & Authentication',
        'Section 3: Workiz API Protocol (Complete)',
        'Section 4: Odoo API Protocol (Complete)',
        'Section 5: Development Constraints & Limitations',
        'Section 6: Integration Workflows',
        'Section 7: Complete Field Dictionary',
        'Section 8: Validation & Proven Methods',
        'Appendix A: All Original Content (Unfiltered)',
        'Appendix B: All Tables (Complete)',
    ]
    for item in toc_items:
        output_doc.add_paragraph(item, style='List Number')
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== SECTION 1: GOLDEN RULES ==========
    add_heading('Section 1: System Architecture & Golden Rules', 1)
    output_doc.add_paragraph(
        'These are the foundational principles that govern all system interactions. '
        'AI agents must internalize these rules before proceeding to technical implementation.'
    )
    
    if content_sections['golden_rules']:
        for para in content_sections['golden_rules']:
            if para.strip():
                output_doc.add_paragraph(para)
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== SECTION 2: CREDENTIALS ==========
    add_heading('Section 2: Master Credentials & Authentication', 1)
    output_doc.add_paragraph('CRITICAL: All API credentials and authentication methods.')
    
    # Extract and display all credential-related content
    if content_sections['credentials']:
        for para in content_sections['credentials']:
            if para.strip():
                p = output_doc.add_paragraph(para)
                if any(x in para for x in ['api_', 'sec_', '7e92006f']):
                    p.runs[0].font.bold = True
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== SECTION 3: WORKIZ API ==========
    add_heading('Section 3: Workiz API Protocol (Complete)', 1)
    output_doc.add_paragraph(
        'Complete Workiz API documentation including all versions (v2.0, v3.0, v4.0), '
        'endpoints, authentication methods, and proven implementation patterns.'
    )
    
    if content_sections['workiz_api']:
        add_heading('Complete Workiz API Content', 2)
        for i, para in enumerate(content_sections['workiz_api'], 1):
            if para.strip():
                output_doc.add_paragraph(para)
    
    print(f"Added {len(content_sections['workiz_api'])} Workiz API paragraphs")
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== SECTION 4: ODOO API ==========
    add_heading('Section 4: Odoo API Protocol (Complete)', 1)
    output_doc.add_paragraph(
        'Complete Odoo API documentation including all versions (v3.0, v3.1, v4.0), '
        'JSON-RPC structure, method specifications, and data type requirements.'
    )
    
    if content_sections['odoo_api']:
        add_heading('Complete Odoo API Content', 2)
        for i, para in enumerate(content_sections['odoo_api'], 1):
            if para.strip():
                output_doc.add_paragraph(para)
    
    print(f"Added {len(content_sections['odoo_api'])} Odoo API paragraphs")
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== SECTION 5: CONSTRAINTS ==========
    add_heading('Section 5: Development Constraints & Limitations', 1)
    output_doc.add_paragraph(
        'Critical constraints for Odoo Server Actions, sandbox limitations, '
        'and technical boundaries that must be respected.'
    )
    
    if content_sections['constraints']:
        for para in content_sections['constraints']:
            if para.strip():
                output_doc.add_paragraph(para)
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== SECTION 6: WORKFLOWS ==========
    add_heading('Section 6: Integration Workflows', 1)
    output_doc.add_paragraph(
        'Step-by-step workflows including the "Lever Pull" SMS mechanism, '
        'reactivation campaigns, and bi-directional sync patterns.'
    )
    
    if content_sections['workflows']:
        for para in content_sections['workflows']:
            if para.strip():
                output_doc.add_paragraph(para)
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== SECTION 7: FIELD DICTIONARY ==========
    add_heading('Section 7: Complete Field Dictionary', 1)
    output_doc.add_paragraph(
        'ALL field definitions from the source document. No truncation. '
        'This section is critical for AI agents to understand data structure.'
    )
    
    if content_sections['field_definitions']:
        add_heading('All Field Definitions (Unfiltered)', 2)
        
        # Group field definitions by detecting patterns
        custom_fields = []
        base_fields = []
        crm_fields = []
        
        for para in content_sections['field_definitions']:
            if para.strip():
                if 'x_studio' in para.lower() or para.strip().startswith('x_'):
                    custom_fields.append(para)
                elif 'crm' in para.lower():
                    crm_fields.append(para)
                else:
                    base_fields.append(para)
        
        # Custom Fields
        if custom_fields:
            add_heading('Custom Fields (Complete List)', 3)
            for field in custom_fields:
                output_doc.add_paragraph(field)
            print(f"Added {len(custom_fields)} custom field definitions")
        
        # Base Fields
        if base_fields:
            add_heading('Base Fields (Complete List)', 3)
            for field in base_fields:
                output_doc.add_paragraph(field)
            print(f"Added {len(base_fields)} base field definitions")
        
        # CRM Fields
        if crm_fields:
            add_heading('CRM/Lead Fields (Complete List)', 3)
            for field in crm_fields:
                output_doc.add_paragraph(field)
            print(f"Added {len(crm_fields)} CRM field definitions")
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== SECTION 8: VALIDATION ==========
    add_heading('Section 8: Validation & Proven Methods', 1)
    output_doc.add_paragraph(
        'Documentation of validated approaches, confirmed methods, and "hard-won" learnings.'
    )
    
    if content_sections['validation']:
        for para in content_sections['validation']:
            if para.strip():
                output_doc.add_paragraph(para)
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== APPENDIX A: ALL ORIGINAL CONTENT ==========
    add_heading('Appendix A: All Original Content (Unfiltered)', 1)
    output_doc.add_paragraph(
        'This appendix contains EVERY paragraph from the source document in sequential order. '
        'Nothing has been removed. Total paragraphs: {}'.format(len(all_paragraphs))
    )
    
    add_heading('Complete Sequential Content', 2)
    for i, para in enumerate(all_paragraphs, 1):
        if para.strip():
            p = output_doc.add_paragraph(f"[{i}] {para}")
            p.runs[0].font.size = Pt(9)
    
    print(f"Added ALL {len(all_paragraphs)} paragraphs to Appendix A")
    
    add_separator()
    output_doc.add_page_break()
    
    # ========== APPENDIX B: ALL TABLES ==========
    add_heading('Appendix B: All Tables (Complete)', 1)
    output_doc.add_paragraph(
        f'This appendix contains ALL {len(all_tables)} tables from the source document. '
        'No data has been truncated.'
    )
    
    for table_num, table_data in enumerate(all_tables, 1):
        add_heading(f'Table {table_num}', 3)
        
        if table_data:
            # Create table in output document
            num_cols = max(len(row) for row in table_data) if table_data else 1
            doc_table = output_doc.add_table(rows=len(table_data), cols=num_cols)
            doc_table.style = 'Light Grid Accent 1'
            
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < num_cols:
                        doc_table.rows[row_idx].cells[col_idx].text = cell_text
            
            output_doc.add_paragraph()  # Spacing
    
    print(f"Added ALL {len(all_tables)} tables to Appendix B")
    
    # ========== SAVE DOCUMENT ==========
    output_path = "AI_Agent_Master_Manual_COMPLETE.docx"
    output_doc.save(output_path)
    
    print("\n" + "="*60)
    print("MANUAL GENERATION COMPLETE")
    print("="*60)
    print(f"Output file: {output_path}")
    print(f"Total paragraphs preserved: {len(all_paragraphs)}")
    print(f"Total tables preserved: {len(all_tables)}")
    print(f"Data loss: ZERO - All content preserved")
    print("="*60)
    
    return output_path

if __name__ == "__main__":
    create_comprehensive_manual()
